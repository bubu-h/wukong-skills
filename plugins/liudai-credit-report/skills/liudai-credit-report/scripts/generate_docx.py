#!/usr/bin/env python3
"""
流贷授信报告 Word 文档生成脚本。

输入：
  - Markdown 格式的报告文本（通过 --markdown 文件或 --text 参数传入）
  - 或 JSON 格式的章节字典（通过 --json 文件传入）
输出：
  - 符合 format-guide.md 格式规范的 .docx 文件

标黄标记：
  ==文本==  包裹的内容自动转换为 Word 黄色底纹（highlight）

依赖：
  pip install python-docx

用法：
  # 从 markdown 文件生成
  python generate_docx.py --markdown report.md -o output.docx

  # 从纯文本传入
  python generate_docx.py --text "报告内容..." -o output.docx

  # 从 JSON 文件生成
  python generate_docx.py --json data.json -o output.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_run_font(run, font_name="宋体"):
    """同时设置西文和东亚字体，防止 Word 回退到 MS 明朝。"""
    run.font.name = font_name
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════════════
#  工具函数
# ═══════════════════════════════════════════════════


def _add_highlight_to_run(run):
    """给 run 添加黄色底纹 highlight。"""
    rpr = run._r.get_or_add_rPr()
    # 移除已有的 highlight
    existing = rpr.find(qn("w:highlight"))
    if existing is not None:
        rpr.remove(existing)
    highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
    rpr.append(highlight)


def _parse_highlight(text):
    """
    解析 ==高亮== 标记，返回 [(文本, 是否高亮), ...] 列表。

    规则：
    - == 成对出现，包裹的文本标黄
    - 支持同一段落内多处标黄
    - 嵌套或不匹配的 == 当作普通文本处理
    """
    parts = re.split(r"(==)", text)
    result = []
    i = 0
    while i < len(parts):
        if parts[i] == "==":
            i += 1
            buf = []
            while i < len(parts) and parts[i] != "==":
                buf.append(parts[i])
                i += 1
            if i < len(parts) and parts[i] == "==":
                i += 1  # 跳过闭合 ==
            content = "".join(buf)
            if content:
                result.append((content, True))
            else:
                pass  # 空的 == 直接丢弃
        else:
            if parts[i]:
                result.append((parts[i], False))
            i += 1
    return result


def _parse_formatting(text):
    """
    解析 **加粗** 和 ==高亮== 标记，返回 [(文本, 是否加粗, 是否高亮), ...] 列表。

    处理顺序：先解析 ==高亮==，再在分段内解析 **加粗**，
    可正确处理嵌套情况如 ==**文本**==。
    """
    # 先解析 ==高亮==
    hl_parts = _parse_highlight(text)

    result = []
    for seg_text, is_hl in hl_parts:
        # 在每一段内解析 **加粗**
        bold_parts = re.split(r"(\*\*[^*]+\*\*)", seg_text)
        for bp in bold_parts:
            if bp.startswith("**") and bp.endswith("**") and len(bp) > 4:
                result.append((bp[2:-2], True, is_hl))
            else:
                result.append((bp, False, is_hl))
    return result


def _set_cell_borders(cell):
    """为单元格设置 0.5pt 黑色实线边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is not None:
        tcPr.remove(borders)
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        "</w:tcBorders>"
    )
    tcPr.append(parse_xml(borders_xml))


def _set_table_borders(table):
    """给表格的所有单元格加边框。"""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


# ═══════════════════════════════════════════════════
#  写入段落
# ═══════════════════════════════════════════════════


def add_paragraph(
    doc,
    text,
    *,
    bold=False,
    font_size=Pt(12),
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=True,
    line_spacing=Pt(20),
):
    """
    添加一个段落，支持 **加粗** 和 ==高亮== 标记。

    参数：
    - doc: Document
    - text: 段落文本（可含 **加粗** 和 ==高亮== 标记）
    - bold: 是否加粗
    - font_size: 字号，默认 12pt
    - alignment: 对齐方式
    - first_line_indent: 是否首行缩进 2 字符
    - line_spacing: 行距，默认固定值 20pt
    """
    if not text.strip():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(20)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return p

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.85) if first_line_indent else Cm(0)
    p.alignment = alignment

    # 解析 **加粗** 和 ==高亮== 并写入 run
    parts = _parse_formatting(text)

    for part_text, is_bold_inline, is_highlighted in parts:
        if not part_text:
            continue
        run = p.add_run(part_text)
        _set_run_font(run)
        run.font.size = font_size
        run.font.bold = bold or is_bold_inline
        if is_highlighted:
            _add_highlight_to_run(run)

    return p


def add_title(doc, text):
    """封面/文档大标题：宋体 18pt 加粗居中。"""
    return add_paragraph(
        doc,
        text,
        bold=True,
        font_size=Pt(18),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )


def add_customer_title(doc, text):
    """客户名称标题：宋体 16pt 加粗居中。"""
    return add_paragraph(
        doc,
        text,
        bold=True,
        font_size=Pt(16),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )


def add_heading(doc, text):
    """章节标题（一、二、三...）：宋体 12pt 加粗首行缩进。"""
    return add_paragraph(
        doc,
        text,
        bold=True,
        font_size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=True,
    )


def add_subheading(doc, text):
    """子标题（1、2、3... / (1)(2)...）：宋体 12pt 加粗首行缩进。"""
    return add_paragraph(
        doc,
        text,
        bold=True,
        font_size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=True,
    )


def add_body(doc, text):
    """正文段落：宋体 12pt 常规首行缩进。"""
    return add_paragraph(
        doc,
        text,
        bold=False,
        font_size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=True,
    )


def add_table_note(doc, text):
    """表注（"单位：万元"）：宋体 11pt 右对齐无缩进。"""
    return add_paragraph(
        doc,
        text,
        bold=False,
        font_size=Pt(11),
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False,
    )


# ═══════════════════════════════════════════════════
#  写入表格
# ═══════════════════════════════════════════════════


def add_table(doc, headers, rows):
    """
    添加带边框的格式化表格。

    参数：
    - headers: [str, ...]
    - rows: [[str, ...], ...]
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True

    # 表头
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.bold = True
        _set_cell_borders(cell)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(value) if value is not None else ""

            # 数字右对齐，文本左对齐
            if re.search(r"^-?\d[\d,\.%]*$", text):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 解析 **加粗** 和 ==高亮==
            cell_parts = _parse_formatting(text)
            for part_text, is_bold_inline, is_hl in cell_parts:
                if not part_text:
                    continue
                run = p.add_run(part_text)
                _set_run_font(run)
                run.font.size = Pt(10.5)
                run.font.bold = is_bold_inline
                if is_hl:
                    _add_highlight_to_run(run)

            _set_cell_borders(cell)

    return table


# ═══════════════════════════════════════════════════
#  解析 Markdown
# ═══════════════════════════════════════════════════


def _is_separator_line(line):
    """
    判断是否为 Markdown 表格分隔行，如 |---|---|。
    分隔行中除 | 外只能包含 -、:、空格。
    """
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    # 去掉首尾 | 后，剩余部分按 | 分割
    inner = stripped[1:-1]
    cells = inner.split("|")
    return all(re.match(r"^\s*:?-+:?\s*$", cell) for cell in cells)


def _is_table_line(line):
    """判断是否为 Markdown 表格数据行（有 | 且不是分隔行）。"""
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    return not _is_separator_line(stripped)


def _try_parse_markdown_table(lines, start):
    """
    尝试从 lines[start] 开始解析一个 Markdown 表格。
    返回 (headers, rows, 消耗的行数)。
    若无法解析返回 (None, None, 0)。
    """
    if start >= len(lines):
        return None, None, 0

    header_line = lines[start].strip()
    if not _is_table_line(header_line):
        return None, None, 0

    headers = [c.strip() for c in header_line.strip("|").split("|")]

    # 第二行必须是分隔行，跳过
    if start + 1 >= len(lines):
        return None, None, 0
    sep_line = lines[start + 1].strip()
    if not _is_separator_line(sep_line):
        return None, None, 0

    # 后续数据行
    rows = []
    i = start + 2
    while i < len(lines) and _is_table_line(lines[i]):
        row_data = [c.strip() for c in lines[i].strip("|").split("|")]
        # 如果某行列数与 header 不一致，丢掉多余或补充空
        while len(row_data) < len(headers):
            row_data.append("")
        row_data = row_data[: len(headers)]
        rows.append(row_data)
        i += 1

    return headers, rows, i - start


def parse_markdown(markdown_text):
    """
    将 Markdown 文本解析为可渲染的元素列表。
    每项为 dict，包含 type 和数据。

    类型:
    - title / customer_title / heading / subheading / body / table_note
    - table: {headers, rows}
    """
    lines = markdown_text.split("\n")
    elements = []
    i = 0
    in_table = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 尝试解析表格
        headers, rows, consumed = _try_parse_markdown_table(lines, i)
        if headers is not None:
            elements.append({"type": "table", "headers": headers, "rows": rows})
            i += consumed
            continue

        # Markdown 标题
        if stripped.startswith("# ") or stripped.startswith("#　"):
            elements.append({"type": "title", "text": re.sub(r"^#\s*", "", stripped)})

        elif stripped.startswith("## ") or stripped.startswith("##　"):
            elements.append({"type": "heading", "text": re.sub(r"^##\s*", "", stripped)})

        elif stripped.startswith("### ") or stripped.startswith("###　"):
            elements.append({"type": "subheading", "text": re.sub(r"^###\s*", "", stripped)})

        elif stripped == "单位：万元":
            elements.append({"type": "table_note", "text": stripped})

        elif re.fullmatch(r"-{3,}", stripped):
            # 三个以上连续的 - 作为分页符
            elements.append({"type": "page_break"})

        else:
            # 正文
            elements.append({"type": "body", "text": stripped})

        i += 1

    return elements


def render_elements(doc, elements):
    """将解析后的 elements 渲染到 doc 中。"""
    for el in elements:
        t = el["type"]
        if t == "title":
            add_title(doc, el["text"])
        elif t == "customer_title":
            add_customer_title(doc, el["text"])
        elif t == "heading":
            add_heading(doc, el["text"])
        elif t == "subheading":
            add_subheading(doc, el["text"])
        elif t == "body":
            add_body(doc, el["text"])
        elif t == "table_note":
            add_table_note(doc, el["text"])
        elif t == "table":
            add_table(doc, el["headers"], el["rows"])
        elif t == "page_break":
            doc.add_page_break()


# ═══════════════════════════════════════════════════
#  主流程
# ═══════════════════════════════════════════════════


def generate_report(content, output_path):
    """
    生成 Word 文档。

    参数：
    - content: str（Markdown 文本）或 dict（章节字典）
    - output_path: 输出文件路径
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)

    # ── 内容处理 ──
    if isinstance(content, str):
        elements = parse_markdown(content)
        render_elements(doc, elements)
    elif isinstance(content, dict):
        # dict 模式：{'sections': {'标题': '内容', ...}, ...}
        if content.get("company_name"):
            add_customer_title(doc, content["company_name"])
        if content.get("report_date"):
            add_body(doc, f"报告日期：{content['report_date']}")

        for section_title, section_content in content.get("sections", {}).items():
            add_heading(doc, section_title)
            # section_content 可以是 str 或 list
            if isinstance(section_content, str):
                add_body(doc, section_content)
            elif isinstance(section_content, list):
                for item in section_content:
                    if isinstance(item, dict):
                        t = item.get("type", "body")
                        if t == "table":
                            add_table(doc, item["headers"], item["rows"])
                        else:
                            add_body(doc, item.get("text", ""))
                    else:
                        add_body(doc, str(item))
    else:
        raise TypeError("content 必须是 str（Markdown）或 dict")

    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="流贷授信报告 Word 文档生成工具",
    )
    input_group = parser.add_mutually_exclusive_group()
    input_group.add_argument(
        "--markdown", "-m", type=str, help="Markdown 文件路径"
    )
    input_group.add_argument(
        "--text", "-t", type=str, help="直接传入 Markdown 文本"
    )
    input_group.add_argument(
        "--json", "-j", type=str, help="JSON 文件路径（章节字典格式）"
    )
    parser.add_argument(
        "--output", "-o", type=str, required=True, help="输出 .docx 文件路径"
    )
    parser.add_argument(
        "--annotated", action="store_true",
        help="同时输出注释版（保留【来源：xxx】标记），文件名为 {output}_注释版.docx"
    )

    args = parser.parse_args()

    if args.markdown:
        raw_content = Path(args.markdown).read_text(encoding="utf-8")
    elif args.text:
        raw_content = args.text
    elif args.json:
        raw_content = json.loads(Path(args.json).read_text(encoding="utf-8"))
    else:
        print("❌ 请指定输入源：--markdown / --text / --json")
        sys.exit(1)

    # 生成清洁版（去除 【来源：xxx】 注释）
    if isinstance(raw_content, str):
        clean_content = re.sub(r"【来源：[^】]+】", "", raw_content)
        has_annotations = "【来源：" in raw_content
        print(f"[DEBUG] 清洁版生成: {'含来源标记' if has_annotations else '无来源标记'}, "
              f"原文长度={len(raw_content)}, 去除后长度={len(clean_content)}")
    else:
        clean_content = raw_content  # dict mode, skip stripping
    generate_report(clean_content, args.output)

    # 生成注释版（保留 【来源：xxx】 注释）
    if args.annotated:
        from pathlib import Path as P
        out = P(args.output)
        annotated_path = str(out.parent / f"{out.stem}_注释版{out.suffix}")
        generate_report(raw_content, annotated_path)
        print(f"📎 注释版：{annotated_path}")


if __name__ == "__main__":
    main()
